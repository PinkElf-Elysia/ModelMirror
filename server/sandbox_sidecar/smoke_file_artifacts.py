"""Contract and real offline artifact smoke for staged Wave 18A file adapters."""

from __future__ import annotations

import argparse
import asyncio
import hashlib
import json
import os
import shutil
import socket
import subprocess
import sys
import tempfile
import threading
import time
from pathlib import Path
from typing import Any

from mcp import ClientSession, StdioServerParameters
from mcp.client.stdio import stdio_client

from .file_artifacts import (
    PANDOC_VERSION,
    WAVE18A_BUILDERS,
    WAVE18A_SCHEMA_SHA256,
    WAVE18A_TOOL_NAMES,
)
from .file_mcp import opaque_file_id


BLOCKED_TOOL = {
    "zcaceres-markdownify-mcp": "webpage-to-markdown",
    "vivekvells-mcp-pandoc": "convert-contents",
    "antvis-mcp-server-chart": "generate_district_map",
}


def _copy_proxy_stdin(sock: socket.socket) -> None:
    try:
        while True:
            chunk = os.read(sys.stdin.fileno(), 64 * 1024)
            if not chunk:
                break
            sock.sendall(chunk)
    except (BrokenPipeError, ConnectionError, OSError):
        pass
    finally:
        try:
            sock.shutdown(socket.SHUT_WR)
        except OSError:
            pass


def _timeout_proxy(adapter_id: str) -> int:
    if adapter_id != "vivekvells-mcp-pandoc":
        return 64
    socket_path = Path(os.environ["MCP_FILES_SOCKET_PATH"])
    workspace_id = os.environ["MCP_FILE_WORKSPACE_ID"]
    sock = socket.socket(socket.AF_UNIX, socket.SOCK_STREAM)
    try:
        sock.settimeout(5)
        sock.connect(str(socket_path))
        sock.sendall(
            json.dumps(
                {
                    "action": "mcp_stdio",
                    "adapter_id": adapter_id,
                    "workspace_id": workspace_id,
                },
                separators=(",", ":"),
            ).encode()
            + b"\n"
        )
        response = json.loads(sock.makefile("rb", buffering=0).readline(4097).decode())
        if response.get("ok") is not True:
            return 69
        sock.settimeout(None)
        threading.Thread(target=_copy_proxy_stdin, args=(sock,), daemon=True).start()
        while True:
            chunk = sock.recv(64 * 1024)
            if not chunk:
                break
            sys.stdout.buffer.write(chunk)
            sys.stdout.buffer.flush()
        return 0
    except (ConnectionError, OSError, ValueError, json.JSONDecodeError):
        return 69
    finally:
        sock.close()


def _digest(tools: list[Any]) -> str:
    reviewed = [
        {"name": tool.name, "inputSchema": tool.inputSchema}
        for tool in sorted(tools, key=lambda item: item.name)
    ]
    return hashlib.sha256(
        json.dumps(
            reviewed,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()


async def contract_only() -> None:
    for adapter_id, builder in sorted(WAVE18A_BUILDERS.items()):
        tools = await builder(object()).list_tools()
        names = {tool.name for tool in tools}
        digest = _digest(tools)
        if names != set(WAVE18A_TOOL_NAMES[adapter_id]):
            raise RuntimeError("wave18a_tool_contract_drift")
        if digest != WAVE18A_SCHEMA_SHA256[adapter_id]:
            raise RuntimeError("wave18a_schema_contract_drift")
        print(
            f"adapter={adapter_id} tools={len(names)} schema_sha256={digest}",
            flush=True,
        )
    version = subprocess.run(
        ["/usr/local/bin/pandoc", "--version"],
        check=True,
        capture_output=True,
        text=True,
        timeout=5,
    ).stdout.splitlines()[0]
    if version != f"pandoc {PANDOC_VERSION}":
        raise RuntimeError("wave18a_pandoc_version_drift")
    print("wave18a_file_contract_smoke=ok", flush=True)


def _base_env(root: Path, workspace_id: str) -> dict[str, str]:
    temp_root = root / "tmp"
    temp_root.mkdir(parents=True, exist_ok=True)
    return {
        "PATH": "/usr/local/bin:/usr/bin:/bin",
        "PYTHONPATH": "/opt/modelmirror",
        "PYTHONDONTWRITEBYTECODE": "1",
        "PYTHONNOUSERSITE": "1",
        "PYTHONUNBUFFERED": "1",
        "HOME": str(temp_root),
        "TMPDIR": str(temp_root),
        "LANG": "C.UTF-8",
        "LC_ALL": "C.UTF-8",
        "TZ": "UTC",
        "NO_PROXY": "*",
        "no_proxy": "*",
        "OPENBLAS_NUM_THREADS": "1",
        "OMP_NUM_THREADS": "1",
        "MKL_NUM_THREADS": "1",
        "NUMEXPR_NUM_THREADS": "1",
        "MCP_FILE_WORKSPACE_ID": workspace_id,
        "MCP_FILE_ADAPTER_ID": "",
        "MCP_FILE_INPUT_ROOT": str(root / "inputs"),
        "MCP_FILE_OUTPUT_ROOT": str(root / "outputs"),
        "MCP_FILE_MEMORY_ROOT": str(root / "memory"),
    }


def _write_fixture(adapter_id: str, input_root: Path) -> tuple[str | None, str]:
    if adapter_id == "zcaceres-markdownify-mcp":
        from docx import Document

        path = input_root / "source.docx"
        document = Document()
        document.add_heading("Wave 18A", level=1)
        document.add_paragraph("offline deterministic Markdownify artifact")
        document.save(path)
        return opaque_file_id(input_root.name, path.name), path.name
    if adapter_id == "vivekvells-mcp-pandoc":
        path = input_root / "source.md"
        path.write_text("# Wave 18A\n\nPandoc deterministic artifact.\n", encoding="utf-8")
        return opaque_file_id(input_root.name, path.name), path.name
    return None, ""


async def _call_ok(session: ClientSession, name: str, arguments: dict[str, Any]) -> Any:
    result = await session.call_tool(name, arguments)
    if result.isError:
        raise RuntimeError(f"wave18a_representative_call_failed:{name}")
    return result


async def _one_round(adapter_id: str, root: Path, round_number: int) -> dict[str, bytes]:
    workspace_id = "mcpws_" + hashlib.sha256(
        f"{adapter_id}:{round_number}".encode()
    ).hexdigest()[:32]
    input_root = root / "inputs" / workspace_id
    output_root = root / "outputs" / workspace_id
    memory_root = root / "memory" / workspace_id
    input_root.mkdir(parents=True)
    output_root.mkdir(parents=True)
    memory_root.mkdir(parents=True)
    file_id, source_name = _write_fixture(adapter_id, input_root)
    source_hash = (
        hashlib.sha256((input_root / source_name).read_bytes()).hexdigest()
        if source_name
        else ""
    )
    env = _base_env(root, workspace_id)
    env["MCP_FILE_ADAPTER_ID"] = adapter_id
    params = StdioServerParameters(
        command=sys.executable,
        args=[
            "-m",
            "sandbox_sidecar.file_landlock_exec",
            "--",
            sys.executable,
            "-m",
            "sandbox_sidecar.file_mcp",
            adapter_id,
        ],
        env=env,
        cwd=Path(env["TMPDIR"]),
    )
    async with stdio_client(params) as (read_stream, write_stream):
        async with ClientSession(read_stream, write_stream) as session:
            await session.initialize()
            listed = await session.list_tools()
            if {tool.name for tool in listed.tools} != set(WAVE18A_TOOL_NAMES[adapter_id]):
                raise RuntimeError("wave18a_runtime_tool_drift")
            if _digest(listed.tools) != WAVE18A_SCHEMA_SHA256[adapter_id]:
                raise RuntimeError("wave18a_runtime_schema_drift")
            if adapter_id == "zcaceres-markdownify-mcp":
                assert file_id is not None
                await _call_ok(
                    session,
                    "docx-to-markdown",
                    {"file_id": file_id, "artifact_name": "document.md"},
                )
                blocked = await session.call_tool(
                    BLOCKED_TOOL[adapter_id], {"url": "https://example.com"}
                )
                if not blocked.isError:
                    raise RuntimeError("wave18a_blocked_network_tool_exposed")
            elif adapter_id == "vivekvells-mcp-pandoc":
                assert file_id is not None
                base = {
                    "file_id": file_id,
                    "input_format": "markdown",
                }
                await _call_ok(
                    session,
                    "convert-contents",
                    {**base, "output_format": "html", "artifact_name": "document.html"},
                )
                await _call_ok(
                    session,
                    "convert-contents",
                    {**base, "output_format": "docx", "artifact_name": "document.docx"},
                )
                blocked = await session.call_tool(
                    BLOCKED_TOOL[adapter_id],
                    {
                        **base,
                        "output_format": "html",
                        "artifact_name": "blocked.html",
                        "filters": ["/tmp/evil"],
                    },
                )
                if not blocked.isError:
                    raise RuntimeError("wave18a_pandoc_filter_exposed")
            else:
                await _call_ok(
                    session,
                    "generate_line_chart",
                    {
                        "data": [
                            {"time": "2025", "value": 4},
                            {"time": "2026", "value": 7},
                        ],
                        "title": "Wave 18A",
                        "artifact_name": "line.png",
                    },
                )
                await _call_ok(
                    session,
                    "generate_bar_chart",
                    {
                        "data": [
                            {"category": "A", "value": 2},
                            {"category": "B", "value": 5},
                        ],
                        "stack": False,
                        "artifact_name": "bar.png",
                    },
                )
                await _call_ok(
                    session,
                    "generate_pie_chart",
                    {
                        "data": [
                            {"category": "A", "value": 2},
                            {"category": "B", "value": 5},
                        ],
                        "artifact_name": "pie.png",
                    },
                )
                blocked = await session.call_tool(BLOCKED_TOOL[adapter_id], {})
                if not blocked.isError:
                    raise RuntimeError("wave18a_remote_chart_tool_exposed")
    if source_name and hashlib.sha256((input_root / source_name).read_bytes()).hexdigest() != source_hash:
        raise RuntimeError("wave18a_source_mutated")
    artifacts = {
        path.name: path.read_bytes()
        for path in sorted(output_root.iterdir())
        if path.is_file() and not path.is_symlink()
    }
    if not artifacts:
        raise RuntimeError("wave18a_artifact_missing")
    for name, data in artifacts.items():
        if name.endswith(".png") and not data.startswith(b"\x89PNG\r\n\x1a\n"):
            raise RuntimeError("wave18a_png_invalid")
        if len(data) > 32 * 1024 * 1024:
            raise RuntimeError("wave18a_artifact_too_large")
    return artifacts


async def _wait_for_socket(socket_path: Path) -> None:
    deadline = time.monotonic() + 10
    while time.monotonic() < deadline:
        if socket_path.is_socket():
            return
        await asyncio.sleep(0.05)
    raise RuntimeError("wave18a_file_server_start_timeout")


async def _stop_process(process: asyncio.subprocess.Process) -> None:
    if process.returncode is not None:
        return
    process.terminate()
    try:
        await asyncio.wait_for(process.wait(), timeout=3)
    except asyncio.TimeoutError:
        process.kill()
        await process.wait()


async def _start_file_server(
    root: Path, *, allowed_adapters: str | None
) -> tuple[asyncio.subprocess.Process, Path]:
    socket_path = root / "run" / "files-mcp.sock"
    socket_path.parent.mkdir(parents=True, exist_ok=True)
    # Reused smoke roots retain the previous Unix-socket pathname after the
    # server process exits.  Remove it before spawning so the readiness probe
    # cannot mistake a stale inode for the new server.
    socket_path.unlink(missing_ok=True)
    workspace_id = "mcpws_" + "f" * 32
    env = _base_env(root, workspace_id)
    env.update(
        {
            "MCP_FILES_SOCKET_PATH": str(socket_path),
            "MCP_FILES_MAX_SESSIONS": "1",
        }
    )
    if allowed_adapters is not None:
        env["MCP_FILE_ALLOWED_ADAPTERS"] = allowed_adapters
    process = await asyncio.create_subprocess_exec(
        sys.executable,
        "-m",
        "sandbox_sidecar.file_server",
        cwd=Path(env["TMPDIR"]),
        env=env,
        stdin=asyncio.subprocess.DEVNULL,
        stdout=asyncio.subprocess.DEVNULL,
        stderr=asyncio.subprocess.DEVNULL,
    )
    try:
        await _wait_for_socket(socket_path)
    except Exception:
        await _stop_process(process)
        raise
    return process, socket_path


async def _exact_allowlist_deny_probe(root: Path) -> None:
    workspace_id = "mcpws_" + "d" * 32
    (root / "inputs" / workspace_id).mkdir(parents=True, exist_ok=True)
    process, socket_path = await _start_file_server(
        root,
        allowed_adapters="markitdown-mcp",
    )
    try:
        reader, writer = await asyncio.open_unix_connection(str(socket_path))
        writer.write(
            json.dumps(
                {
                    "action": "mcp_stdio",
                    "adapter_id": "vivekvells-mcp-pandoc",
                    "workspace_id": workspace_id,
                },
                separators=(",", ":"),
            ).encode()
            + b"\n"
        )
        await writer.drain()
        response = json.loads((await asyncio.wait_for(reader.readline(), timeout=3)).decode())
        writer.close()
        await writer.wait_closed()
        if response.get("ok") is not False or response.get("code") != "mcp_adapter_denied":
            raise RuntimeError("wave18a_exact_allowlist_deny_failed")
    finally:
        await _stop_process(process)


def _pandoc_process_count() -> int:
    count = 0
    for entry in Path("/proc").iterdir():
        if not entry.name.isdigit():
            continue
        try:
            command = (entry / "cmdline").read_bytes()
        except OSError:
            continue
        if b"/usr/local/bin/pandoc\0" in command:
            count += 1
    return count


async def _timeout_probe(root: Path) -> None:
    adapter_id = "vivekvells-mcp-pandoc"
    workspace_id = "mcpws_" + "e" * 32
    input_root = root / "inputs" / workspace_id
    input_root.mkdir(parents=True, exist_ok=True)
    source = input_root / "timeout.md"
    source.write_bytes(b"# Timeout\n\n" + b"safe paragraph\n" * 800_000)
    file_id = opaque_file_id(workspace_id, source.name)
    process, socket_path = await _start_file_server(
        root, allowed_adapters=adapter_id
    )
    proxy_env = _base_env(root, workspace_id)
    proxy_env["MCP_FILES_SOCKET_PATH"] = str(socket_path)
    params = StdioServerParameters(
        command=sys.executable,
        args=["-m", "sandbox_sidecar.smoke_file_artifacts", "--proxy", adapter_id],
        env=proxy_env,
        cwd=Path(proxy_env["TMPDIR"]),
    )
    started = time.monotonic()
    try:
        async with stdio_client(params) as (read_stream, write_stream):
            async with ClientSession(read_stream, write_stream) as session:
                await session.initialize()
                try:
                    await asyncio.wait_for(
                        session.call_tool(
                            "convert-contents",
                            {
                                "file_id": file_id,
                                "input_format": "markdown",
                                "output_format": "docx",
                                "artifact_name": "timeout.docx",
                            },
                        ),
                        timeout=0.05,
                    )
                except asyncio.TimeoutError:
                    pass
                else:
                    raise RuntimeError("wave18a_timeout_not_triggered")
        deadline = time.monotonic() + 3
        while time.monotonic() < deadline and _pandoc_process_count():
            await asyncio.sleep(0.05)
        if _pandoc_process_count():
            raise RuntimeError("wave18a_timeout_process_leaked")
        if time.monotonic() - started > 5:
            raise RuntimeError("wave18a_timeout_cleanup_slow")
    finally:
        await _stop_process(process)


async def runtime_smoke() -> None:
    base = Path(tempfile.mkdtemp(prefix="wave18a-file-smoke-"))
    try:
        await _exact_allowlist_deny_probe(base / "exact-allowlist-deny")
        first: dict[str, dict[str, bytes]] = {}
        for adapter_id in sorted(WAVE18A_BUILDERS):
            first[adapter_id] = await asyncio.wait_for(
                _one_round(adapter_id, base / "round-1", 1), timeout=55
            )
            second = await asyncio.wait_for(
                _one_round(adapter_id, base / "round-2", 2), timeout=55
            )
            if first[adapter_id] != second:
                raise RuntimeError("wave18a_artifact_not_deterministic")
            print(
                f"adapter={adapter_id} rounds=2 artifacts={len(second)} "
                f"sha256={','.join(hashlib.sha256(data).hexdigest() for data in second.values())}",
                flush=True,
            )
        await _timeout_probe(base / "timeout")
    finally:
        shutil.rmtree(base, ignore_errors=False)
    if base.exists():
        raise RuntimeError("wave18a_cleanup_failed")
    print(
        "wave18a_file_runtime_smoke=ok network=none source_immutable=true "
        "deterministic=true exact_allowlist_deny=true timeout=verified cleanup=verified",
        flush=True,
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--contract-only", action="store_true")
    parser.add_argument("--proxy", choices=("vivekvells-mcp-pandoc",))
    args = parser.parse_args()
    if args.proxy:
        raise SystemExit(_timeout_proxy(args.proxy))
    asyncio.run(contract_only() if args.contract_only else runtime_smoke())


if __name__ == "__main__":
    main()
